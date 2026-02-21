# /// script
# dependencies = [
#   "python-docx",
# ]
# ///

from docx import Document
from docx.shared import Pt, Inches
import os

os.makedirs('sample_docs', exist_ok=True)

# 1. Purchase Order
doc1 = Document()
doc1.add_heading('PURCHASE ORDER', 0)
doc1.add_paragraph('PO Number: PO-PHARMA-2026')
doc1.add_paragraph('Date: October 24, 2026')
doc1.add_paragraph('Buyer: Global Health Distributors\n123 MedSupply Way, New York, NY 10001')
doc1.add_paragraph('Supplier: BioGen Pharmaceuticals\n456 Science Park, Boston, MA 02115')

doc1.add_heading('Items Ordered', level=1)
table = doc1.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Quantity'
hdr_cells[2].text = 'Unit Price'
hdr_cells[3].text = 'Total'

items = [
    ('Insulin Glargine 100U/mL (Vials)', '5000', '$25.00', '$125,000.00'),
    ('Amoxicillin 500mg (Bottles of 100)', '2000', '$12.00', '$24,000.00'),
    ('Lisinopril 10mg (Bottles of 90)', '3000', '$8.50', '$25,500.00')
]

for item, qty, price, total in items:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = qty
    row_cells[2].text = price
    row_cells[3].text = total

doc1.add_paragraph('\nTotal Amount: $174,500.00').bold = True
doc1.add_paragraph('Special Instructions: Items must be kept refrigerated at 2°C to 8°C during transit (Strict Cold Chain).')

doc1.save('sample_docs/Purchase_Order_Pharma.docx')

# 2. Invoice
doc2 = Document()
doc2.add_heading('COMMERCIAL INVOICE', 0)
doc2.add_paragraph('Invoice Number: INV-PHARMA-2026-09')
doc2.add_paragraph('Date: October 25, 2026')
doc2.add_paragraph('PO Reference: PO-PHARMA-2026')
doc2.add_paragraph('Bill To: Global Health Distributors\n123 MedSupply Way, New York, NY 10001')
doc2.add_paragraph('Make Checks Payable To: BioGen Pharmaceuticals\n456 Science Park, Boston, MA 02115')

doc2.add_heading('Charges', level=1)
table2 = doc2.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2_cells = table2.rows[0].cells
hdr2_cells[0].text = 'Description'
hdr2_cells[1].text = 'Quantity'
hdr2_cells[2].text = 'Unit Price'
hdr2_cells[3].text = 'Amount'

for item, qty, price, total in items:
    row_cells = table2.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = qty
    row_cells[2].text = price
    row_cells[3].text = total

doc2.add_paragraph('\nSubtotal: $174,500.00')
doc2.add_paragraph('Shipping & Handling (Cold Chain Express): $4,200.00')
doc2.add_paragraph('Total Due: $178,700.00').bold = True
doc2.add_paragraph('Payment Terms: Net 30')

doc2.save('sample_docs/Invoice_Pharma.docx')

# 3. Bill of Lading
doc3 = Document()
doc3.add_heading('BILL OF LADING (BOL)', 0)
doc3.add_paragraph('BOL Number: BOL-PHARMA-8899')
doc3.add_paragraph('Date: October 26, 2026')
doc3.add_paragraph('Shipper: BioGen Pharmaceuticals\n456 Science Park, Boston, MA 02115')
doc3.add_paragraph('Consignee: Global Health Distributors\n123 MedSupply Way, New York, NY 10001')
doc3.add_paragraph('Carrier: ColdChain Logistics Inc.')

doc3.add_heading('Shipment Details', level=1)
table3 = doc3.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr3_cells = table3.rows[0].cells
hdr3_cells[0].text = 'Packages'
hdr3_cells[1].text = 'Description of Articles'
hdr3_cells[2].text = 'Weight (kg)'
hdr3_cells[3].text = 'Hazard Class'

bol_items = [
    ('50 Pallets', 'Insulin Glargine 100U/mL (Refrigerated)', '1200 kg', 'N/A (Temp Controlled)'),
    ('20 Pallets', 'Amoxicillin 500mg', '800 kg', 'N/A'),
    ('30 Pallets', 'Lisinopril 10mg', '950 kg', 'N/A')
]

for pkgs, desc, weight, haz in bol_items:
    row_cells = table3.add_row().cells
    row_cells[0].text = pkgs
    row_cells[1].text = desc
    row_cells[2].text = weight
    row_cells[3].text = haz

doc3.add_paragraph('\nTotal Weight: 2,950 kg').bold = True
doc3.add_paragraph('Special Instructions: \n- MUST MAINTAIN 2°C - 8°C AT ALL TIMES.\n- DO NOT FREEZE.\n- Handle with care.')
doc3.add_paragraph('\nShipper Signature: _______________________')
doc3.add_paragraph('Carrier Signature: _______________________')
doc3.add_paragraph('Consignee Signature: _______________________')

doc3.save('sample_docs/Bill_of_Lading_Pharma.docx')

print("Successfully generated pharmaceutical documents in 'sample_docs/' directory.")
